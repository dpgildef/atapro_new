import streamlit as st
import google.generativeai as genai
import tempfile
import os
import time
from io import BytesIO
from docx import Document 
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH # Importante para justificar texto

# --- 1. CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="AtaPro.pt | Condomínios",
    page_icon="🏢",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# --- ESTILO CSS ---
st.markdown("""
    <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        
        .stButton>button {
            width: 100%;
            border-radius: 8px;
            height: 3.5em;
            font-weight: 600;
            font-size: 16px;
        }

        /* TRUQUE CSS PARA TRADUZIR O UPLOAD */
        [data-testid='stFileUploaderDropzone'] div div span {display: none;}
        [data-testid='stFileUploaderDropzone'] div div::after {
           content: "Arraste e largue os ficheiros aqui";
           font-size: 1.2em; font-weight: bold;
        }
        [data-testid='stFileUploaderDropzone'] small {display: none;}
        /* VISIBILIDADE REFORÇADA NO CSS */
        [data-testid='stFileUploaderDropzone']::after {
           content: "⚠️ Apenas MP3 ou M4A • Máx: 200MB por ficheiro";
           font-size: 1em; 
           font-weight: bold;
           display: block; 
           text-align: center; 
           margin-top: 10px; 
           color: #d63031;
           background-color: #ffeaea;
           padding: 5px;
           border-radius: 5px;
        }
        
        /* Caixa de Aviso Legal */
        .legal-box {
            font-size: 0.85em;
            background-color: #fdf2f2;
            border-left: 4px solid #dc3545;
            padding: 12px;
            margin-top: 10px;
            color: #5a1e1e;
        }
    </style>
""", unsafe_allow_html=True)

# --- 2. SISTEMA DE LOGIN ---
def check_password():
    if st.session_state.get("password_correct", False):
        return st.session_state.get("user_name", "Utilizador")

    st.markdown("<br><br>", unsafe_allow_html=True)
    st.title("🏢 AtaPro | Condomínios")
    st.info("Área exclusiva para Administradores de Condomínio.")
    
    password_input = st.text_input("Introduza a sua Chave de Acesso:", type="password")
    
    if st.button("🔓 Entrar", type="primary"):
        try:
            passwords = st.secrets["passwords"]
            senha_para_nome = {v: k for k, v in passwords.items()}
            
            if password_input in senha_para_nome:
                st.session_state["password_correct"] = True
                st.session_state["user_name"] = senha_para_nome[password_input]
                st.toast(f"Bem-vindo(a), {st.session_state['user_name']}!", icon="👋")
                time.sleep(1)
                st.rerun()
            else:
                st.error("❌ Chave de acesso incorreta.")
        except KeyError:
            st.error("Erro de configuração interna (Secrets).")
    return None

nome_utilizador = check_password()
if not nome_utilizador:
    st.stop()

# ==========================================
# LÓGICA DE ESTADO (MEMÓRIA)
# ==========================================
if "texto_ata_final" not in st.session_state:
    st.session_state["texto_ata_final"] = None

# --- 3. CONFIGURAÇÃO API ---
try:
    API_KEY = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=API_KEY)
except Exception:
    st.error("⚠️ ERRO CRÍTICO: Chave de API Google em falta.")
    st.stop()

# --- 4. FUNÇÃO GERADORA DE WORD (.docx) AVANÇADA ---
def criar_word(texto_ata):
    doc = Document()
    
    # Estilo Global
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Título Principal
    heading = doc.add_heading('ATA DA ASSEMBLEIA DE CONDOMÍNIO', 0)
    heading.alignment = 1 # Centro
    
    # Processamento Inteligente do Texto (Markdown -> Word)
    for paragrafo in texto_ata.split('\n'):
        if paragrafo.strip():
            # Cria parágrafo
            p = doc.add_paragraph()
            # Justificar texto (Padrão Profissional)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Lógica para detetar **negrito**
            # Divide o texto pelos asteriscos
            partes = paragrafo.split('**')
            
            for i, parte in enumerate(partes):
                run = p.add_run(parte)
                # Se o índice for ímpar (1, 3, 5...), é a parte que estava entre **
                if i % 2 != 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0) # Preto forte

    # --- BLOCO DE ASSINATURAS (NOVO) ---
    doc.add_paragraph("\n" * 2) # Espaço
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    
    # Assinatura Presidente
    cell1 = table.cell(0, 0)
    p1 = cell1.paragraphs[0]
    p1.add_run("__________________________\nO Presidente da Mesa").bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Assinatura Secretário
    cell2 = table.cell(0, 1)
    p2 = cell2.paragraphs[0]
    p2.add_run("__________________________\nO Secretário").bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- RODAPÉ DE RESPONSABILIDADE ---
    doc.add_paragraph("_" * 50)
    legal_note = doc.add_paragraph()
    
    texto_disclaimer = (
        "NOTA DE RESPONSABILIDADE: O presente documento constitui uma minuta de auxílio redigida por Inteligência Artificial. "
        "A sua validade legal e eficácia como título executivo (nos termos do DL n.º 268/94 e Lei n.º 8/2022) "
        "dependem obrigatoriamente da revisão, aprovação e assinatura pela Mesa da Assembleia e Administração do Condomínio."
    )
    
    run = legal_note.add_run(texto_disclaimer)
    run.font.size = Pt(7)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 5. PROCESSAMENTO DE ÁUDIO (IA) ---
def processar_ata(files):
    status = st.status("⚙️ A processar ata de condomínio...", expanded=True)
    arquivos_temp = []
    arquivos_gemini = []
    
    try:
        files.sort(key=lambda x: x.name)
        
        # A: Upload
        status.write(f"📤 A carregar {len(files)} ficheiro(s)...")
        for file in files:
            suffix = os.path.splitext(file.name)[1].lower()
            if not suffix: suffix = ".mp3"
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(file.getvalue())
                tmp_path = tmp.name
            
            g_file = genai.upload_file(tmp_path)
            arquivos_gemini.append(g_file)
            arquivos_temp.append(tmp_path) 
            status.write(f"✅ Recebido: {file.name}")

        # B: Espera
        status.write("🎧 A analisar as deliberações dos condóminos...")
        for g_file in arquivos_gemini:
            while g_file.state.name == "PROCESSING":
                time.sleep(2)
                g_file = genai.get_file(g_file.name)
            if g_file.state.name == "FAILED":
                raise Exception(f"Erro ao ler o ficheiro {g_file.name}.")

        # C: Geração
        status.write("✍️ A redigir a ata segundo a Lei n.º 8/2022...")
        model = genai.GenerativeModel("models/gemini-2.5-flash")
        
        # --- PROMPT REFINADO PARA "ATA PERFEITA" ---
        prompt = """
        Tu és um Administrador de Condomínios e Jurista em Portugal. 
        A tua tarefa é redigir uma ATA DE ASSEMBLEIA DE CONDOMÍNIO IMPECÁVEL.

        REGRAS DE FORMATAÇÃO (Obrigatório):
        - Usa **negrito** (entre dois asteriscos) para destacar Títulos, Votações e Valores Monetários.
        - Não uses tabelas nem Markdown complexo (#), apenas texto limpo e negritos.

        ESTRUTURA OBRIGATÓRIA DA ATA:
        1. CABEÇALHO COMPLETO:
           - Número da Ata, Data, Hora Início/Fim, Local exato.
           - Tipo de Reunião (Ordinária/Extraordinária).
           - Identificação do Presidente e Secretário.

        2. PRESENÇAS E QUÓRUM (Crítico):
           - Lista as frações presentes.
           - **Calcula ou menciona a Permilagem/Percentagem total representada.**
           - Declara explicitamente se "existe quórum constitutivo para deliberar".

        3. ORDEM DE TRABALHOS:
           - Copia exatamente os pontos discutidos.

        4. DELIBERAÇÕES (Para cada ponto da Ordem de Trabalhos):
           - **Discussão:** Resumo breve e impessoal.
           - **Votação:** Discrimina claramente: "Votos a Favor (X permilagem)", "Contra (Fração Y)", "Abstenções".
           - **Decisão:** Escreve em letras garrafais: **APROVADO POR UNANIMIDADE** ou **MAIORIA**.
           - **Título Executivo:** Se houver dívidas ou valores aprovados, especifica o valor exato (ex: "1.200,00€") e prazos de pagamento para que a ata sirva de título executivo.

        5. ENCERRAMENTO:
           - "Nada mais havendo a tratar, deu-se por encerrada a sessão às [Hora]..."
        
        Escreve em Português de Portugal (PT-PT) formal e jurídico.
        """
        
        response = model.generate_content([prompt] + arquivos_gemini)
        texto_gerado = response.text
        
        status.update(label="✅ Ata de Condomínio Gerada!", state="complete", expanded=False)
        
        # D: Limpeza
        for g_file in arquivos_gemini:
            try: genai.delete_file(g_file.name)
            except: pass
        for path in arquivos_temp:
            try: os.remove(path)
            except: pass
            
        return texto_gerado

    except Exception as e:
        status.update(label="❌ Erro no processamento", state="error")
        st.error(f"Ocorreu um erro: {e}")
        return None

# --- 6. INTERFACE DE UTILIZADOR ---

col1, col2 = st.columns([3, 1])
with col1:
    st.title("🏢 AtaPro | Condomínios")
    st.markdown(f"**Bem-vindo, {nome_utilizador}.**")
with col2:
    if st.button("Sair 🔒"):
        st.session_state["password_correct"] = False
        st.session_state["user_name"] = None
        st.rerun()

# --- MOSTRAR RESULTADO SE JÁ EXISTIR NA MEMÓRIA ---
if st.session_state["texto_ata_final"]:
    st.success("✅ A sua minuta de ata está pronta.")
    
    # AVISO DE RESPONSABILIDADE
    st.markdown("""
    <div class="legal-box">
    ⚖️ <strong>Aviso de Responsabilidade:</strong><br> 
    Este ficheiro é uma <strong>minuta de trabalho</strong> gerada por IA. <br>
    A AtaPro.pt fornece o serviço de transcrição e redação, mas a <strong>conferência do conteúdo e a validação jurídica final</strong> 
    são da exclusiva responsabilidade da Administração do Condomínio, mediante a aposição das assinaturas legais.
    </div>
    """, unsafe_allow_html=True)
    
    st.write("### 📥 Descarregar Documento")
    
    # Converter para Word
    word_file = criar_word(st.session_state["texto_ata_final"])
    
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        st.download_button(
            label="📄 Descarregar em WORD (.docx)",
            data=word_file,
            file_name="Ata_Condominio.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
    with col_d2:
        if st.button("🔄 Começar Nova Ata"):
            st.session_state["texto_ata_final"] = None
            st.rerun()
            
    with st.expander("👁️ Ver Texto da Ata (Pré-visualização)"):
        st.markdown(st.session_state["texto_ata_final"])

# --- MOSTRAR UPLOAD SE AINDA NÃO HOUVER ATA ---
else:
    with st.expander("🎙️ GUIA: COMO GRAVAR ASSEMBLEIA DE CONDOMÍNIO", expanded=False):
        st.markdown("""
        Para validade legal (Lei n.º 8/2022), comece a gravação dizendo:
        1.  **"Assembleia do Condomínio do prédio sito em [Morada]..."**
        2.  **"Hoje é dia [X], hora [Y]."**
        3.  **"Presenças: Fração A (Sr. João), Fração B (Sra. Maria)..."**
        4.  **"Ponto 1 da Ordem de Trabalhos: [Assunto]..."**
        """)

    st.write("### 1. Carregar Gravações")

    with st.expander("📱 Ajuda para iPhone/WhatsApp"):
        st.info("No iPhone ou WhatsApp: 'Partilhar' > 'Guardar em Ficheiros'.")
    
    # --- AVISO VISÍVEL SOBRE FICHEIROS ---
    st.info("⚠️ **IMPORTANTE:** Carregue apenas ficheiros **MP3** ou **M4A**. Outros formatos (como WAV ou PDF) serão rejeitados para garantir rapidez. Limite: 200MB.")

    uploaded_files = st.file_uploader(
        "Selecione os ficheiros de áudio:", 
        accept_multiple_files=True
    )

    if uploaded_files:
        # --- VERIFICAÇÃO DE FICHEIROS INCORRETOS ---
        ficheiros_validos = True
        for ficheiro in uploaded_files:
            ext = os.path.splitext(ficheiro.name)[1].lower()
            if ext not in ['.mp3', '.m4a']:
                st.error(f"❌ ERRO: O ficheiro '{ficheiro.name}' tem um formato inválido ({ext}).")
                st.error("Por favor, carregue apenas ficheiros **.mp3** ou **.m4a**.")
                ficheiros_validos = False
        
        if ficheiros_validos:
            st.caption(f"📂 {len(uploaded_files)} ficheiro(s) válido(s).")
            
            st.markdown("---")
            st.subheader("🛡️ Privacidade e Termos")
            
            st.markdown("""
            <div>
            <ul>
                <li><strong>Segurança:</strong> Áudio processado via Google Enterprise (encriptado).</li>
                <li><strong>Eliminação:</strong> Dados apagados imediatamente após a geração.</li>
                <li><strong>Sem Histórico:</strong> Não guardamos cópias das atas.</li>
            </ul>
            </div>
            """, unsafe_allow_html=True)
            
            autorizacao = st.checkbox("Li e aceito a Política de Privacidade.")

            if autorizacao:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("📝 GERAR ATA DE CONDOMÍNIO", type="primary"):
                    resultado = processar_ata(uploaded_files)
                    if resultado:
                        st.session_state["texto_ata_final"] = resultado
                        st.rerun() 
            else:
                st.info("👆 Aceite os termos para continuar.")
